"""Structured document loaders — DOCX and XLSX."""
from __future__ import annotations
from groundguard.models.result import Source


def load_docx(path: str, split_by: str = "heading",
              source_type: str = "document_section",
              populate_context: bool = False) -> list[Source]:
    """
    Load a .docx file and split into Source objects.
    split_by="heading": split on Heading 1/2/3 styles
    split_by="paragraph": split on non-empty paragraphs
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError("python-docx required: pip install 'groundguard[loaders]'")

    doc = Document(path)
    sources = []

    if split_by == "heading":
        current_chunks = []
        section_idx = 0
        in_section = False

        for para in doc.paragraphs:
            if para.style.name.startswith("Heading"):
                # Save previous section
                if current_chunks:
                    content = "\n".join(p.text.strip() for p in current_chunks if p.text.strip())
                    if content:
                        sources.append(Source(
                            content=content,
                            source_id=f"section_{section_idx}",
                            source_type=source_type,
                        ))
                        section_idx += 1
                current_chunks = [para]
                in_section = True
            else:
                if in_section or para.text.strip():
                    current_chunks.append(para)
                    if not in_section:
                        in_section = True

        # Flush last section
        if current_chunks:
            content = "\n".join(p.text.strip() for p in current_chunks if p.text.strip())
            if content:
                sources.append(Source(
                    content=content,
                    source_id=f"section_{section_idx}",
                    source_type=source_type,
                ))

    elif split_by == "paragraph":
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                sources.append(Source(
                    content=para.text.strip(),
                    source_id=f"para_{i}",
                    source_type=source_type,
                ))

    return sources


def load_xlsx(path: str, source_type: str = "table_row") -> list[Source]:
    """
    Load an .xlsx file and convert each data row to a Source.
    content format: "Header: value\\nHeader2: value2" (NOT bare numeric)
    """
    try:
        from openpyxl import load_workbook
    except ImportError:
        raise ImportError("openpyxl required: pip install 'groundguard[loaders]'")

    wb = load_workbook(path)
    ws = wb.active
    rows = list(ws.iter_rows(values_only=True))

    if not rows:
        return []

    headers = [str(h) if h is not None else f"col_{i}" for i, h in enumerate(rows[0])]
    sources = []

    for row_idx, row in enumerate(rows[1:], start=1):
        parts = []
        for header, value in zip(headers, row):
            if value is not None:
                parts.append(f"{header}: {value}")
        if parts:
            sources.append(Source(
                content="\n".join(parts),
                source_id=f"row_{row_idx}",
                source_type=source_type,
            ))

    return sources
